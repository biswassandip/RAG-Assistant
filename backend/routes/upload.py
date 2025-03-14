from fastapi import APIRouter, UploadFile, File, HTTPException
import os
import shutil
from datetime import datetime
import pytesseract
import cv2
import numpy as np
import pandas as pd
from bs4 import BeautifulSoup
from langchain_core.documents import Document
from services.vectorstore import vector_store
from constants import URL_UPLOAD, URL_GET_FILES
from PyPDF2 import PdfReader
import docx

router = APIRouter()
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

uploaded_files_metadata = []

# ---------- FILE TEXT EXTRACTION FUNCTIONS ----------


def extract_text(file_path, file_type):
    """Extract text from various file types: PDF, DOCX, TXT, CSV, XML, Images."""
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".pdf":
            with open(file_path, "rb") as f:
                reader = PdfReader(f)
                text = "\n".join([page.extract_text()
                                 for page in reader.pages if page.extract_text()])

        elif ext == ".docx":
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()

        elif ext == ".csv":
            df = pd.read_csv(file_path)
            text = df.head(5).to_string()

        elif ext == ".xml":
            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f, "xml")
                text = soup.get_text()

        elif ext in [".jpg", ".jpeg", ".png", ".bmp", ".tiff"]:
            image = cv2.imread(file_path)
            text = pytesseract.image_to_string(image)
            text = text if text.strip() else "No text detected"

        else:
            text = "Unsupported file type"

        return text if text else "Could not extract text"

    except Exception as e:
        return f"Error extracting text: {str(e)}"

# ---------- FILE UPLOAD HANDLING ----------


@router.post(URL_UPLOAD)
async def upload_files(files: list[UploadFile] = File(...)):
    """Handle file uploads and store metadata in FAISS with improved document handling."""
    global uploaded_files_metadata
    documents = []

    for file in files:
        file_path = os.path.join(UPLOAD_DIR, file.filename)

        # Save the file
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Extract text
        extracted_text = extract_text(file_path, file.content_type)
        upload_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        # Metadata details
        metadata = {
            "uploaded_date": upload_date,
            "file_name": file.filename,
            "file_type": file.content_type or "Unknown",
            "word_count": len(extracted_text.split()) if extracted_text else 0,
            "char_count": len(extracted_text) if extracted_text else 0,
            "summary": extracted_text[:300] + "..." if extracted_text else "No content extracted"
        }

        uploaded_files_metadata.append(metadata)

        # Create a LangChain Document object
        if extracted_text:
            doc = Document(page_content=extracted_text, metadata=metadata)
            documents.append(doc)

    # Add extracted text to FAISS Vector Store
    if documents:
        vector_store.add_documents(documents)

    return {"message": "Files uploaded and processed successfully!", "metadata": uploaded_files_metadata}

# ---------- GET UPLOADED FILES ----------


@router.get(URL_GET_FILES)
def get_uploaded_files():
    """Retrieve uploaded file metadata."""
    return uploaded_files_metadata
